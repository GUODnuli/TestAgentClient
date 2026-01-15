"""
Word文档解析工具

负责从Word文档中提取文本内容、表格和结构化信息。
支持.docx格式（使用python-docx库）。
"""

from typing import Dict, List, Any
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


class WordDocumentParser:
    """
    Word文档解析工具类
    
    功能：
    - 提取纯文本内容
    - 提取结构化内容（保留标题层级、段落、表格）
    - 提取表格数据
    """
    
    def __init__(self):
        pass
    
    def extract_text(self, file_path: str) -> str:
        """
        提取Word文档中的纯文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            提取的纯文本内容
        """
        doc = Document(file_path)
        
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_parts.append(table_text)
        
        return "\n\n".join(text_parts)
    
    def extract_structured_content(self, file_path: str) -> Dict[str, Any]:
        """
        提取结构化内容（段落、表格、标题层级等）
        
        Args:
            file_path: Word文档路径
            
        Returns:
            包含结构化内容的字典
        """
        doc = Document(file_path)
        
        content = {
            "paragraphs": [],
            "tables": [],
            "headings": [],
            "metadata": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        }
        
        # 提取段落和标题
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            
            para_info = {
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal"
            }
            
            # 识别标题
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                heading_level = paragraph.style.name.replace("Heading ", "")
                para_info["is_heading"] = True
                para_info["heading_level"] = heading_level
                content["headings"].append(para_info)
            
            content["paragraphs"].append(para_info)
        
        # 提取表格
        for idx, table in enumerate(doc.tables):
            table_data = self._extract_table_data(table)
            content["tables"].append({
                "index": idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": table_data
            })
        
        return content
    
    def extract_tables(self, file_path: str) -> List[List[str]]:
        """
        提取文档中的所有表格数据
        
        Args:
            file_path: Word文档路径
            
        Returns:
            表格数组，每个表格是一个二维数组
        """
        doc = Document(file_path)
        
        all_tables = []
        
        for table in doc.tables:
            table_data = self._extract_table_data(table)
            if table_data:
                all_tables.append(table_data)
        
        return all_tables
    
    def _extract_table_text(self, table: Table) -> str:
        """提取表格的文本内容（用于纯文本提取）"""
        rows_text = []
        
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            
            if cells_text:
                rows_text.append(" | ".join(cells_text))
        
        return "\n".join(rows_text)
    
    def _extract_table_data(self, table: Table) -> List[List[str]]:
        """提取表格数据为二维数组"""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        
        return table_data
    
    def validate_file(self, file_path: str) -> bool:
        """
        验证文件是否为有效的Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否为有效的.docx文件
        """
        path = Path(file_path)
        
        # 检查文件是否存在
        if not path.exists():
            return False
        
        # 检查扩展名
        if path.suffix.lower() != ".docx":
            return False
        
        # 尝试打开文档
        try:
            Document(file_path)
            return True
        except Exception:
            return False
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        获取文档基本信息
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档信息字典
        """
        doc = Document(file_path)
        path = Path(file_path)
        
        # 获取核心属性
        core_props = doc.core_properties
        
        info = {
            "filename": path.name,
            "file_size_bytes": path.stat().st_size,
            "file_size_kb": round(path.stat().st_size / 1024, 2),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "title": core_props.title if core_props.title else "",
            "author": core_props.author if core_props.author else "",
            "subject": core_props.subject if core_props.subject else "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        return info
